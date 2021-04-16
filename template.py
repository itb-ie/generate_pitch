# this is the template for the pitch
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm

import get_papers


def set_col_widths(table):
    widths = (Cm(3), Cm(13))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def fill_in(table, article):
    columns = table.columns[0]
    columns.cells[0].text = "FOUR"
    columns.cells[1].text = "(A) Full Reference"
    columns.cells[2].text = "(B) Basic Research Question"
    columns.cells[3].text = "(C) Key Paper(s)"
    columns.cells[4].text = "(D) Motivation/Puzzle"
    columns.cells[5].text = "THREE"
    columns.cells[6].text = "(E) Idea?"
    columns.cells[7].text = "(F) Data?"
    columns.cells[8].text = "(G) Tools?"
    columns.cells[9].text = "TWO"
    columns.cells[10].text = "(H) What's New?"
    columns.cells[11].text = "(I) So What?"
    columns.cells[12].text = "ONE"
    columns.cells[13].text = "(J) Contribution?"
    columns.cells[14].text = "+ ONE"
    columns.cells[15].text = "(K) Three Key Findings"

    # this is where my text goes in

    columns = table.columns[1]
    columns.cells[0].text = "Four elements of framing"
    columns.cells[1].text = "{}".format(article["full_reference"])
    columns.cells[2].text = f"{article['answer_research_question']}"
    references = ""
    for idx, ref in enumerate(article["top_references"]):
        references += f"{idx+1}. {ref['auth']}, {ref['year']}, {ref['rest']}.\n"
    columns.cells[3].text = "{}".format(references)
    columns.cells[4].text = "{}".format(article["answer_motivation"])
    columns.cells[5].text = "Three building blocks"
    columns.cells[6].text = f"{article['answer_idea']}"
    columns.cells[7].text = "{}".format("\n".join(article["answer_data"]))
    columns.cells[8].text = f"{article['answer_tools']}"
    columns.cells[9].text = "Two key questions"
    columns.cells[10].text = f"{article['answer_whats_new']}"
    columns.cells[11].text = f"{article['answer_so_what']}"
    columns.cells[12].text = "One bottom line"
    columns.cells[13].text = "{}".format(article["answer_contribution"])
    columns.cells[14].text = ""
    columns.cells[15].text = "{}".format(article["answer_key_findings"])


if __name__ == "__main__":
    doc = docx.Document()
    styles = doc.styles
    paragraph_styles = [
        s for s in styles if s.type == WD_STYLE_TYPE.TABLE
    ]
    for style in paragraph_styles:
        print(style.name)

    doc = docx.Document()
    table = doc.add_table(16, 2, style="Light Grid Accent 1")
    set_col_widths(table)
    fill_in(table)

    doc.save("Pitch.docx")


